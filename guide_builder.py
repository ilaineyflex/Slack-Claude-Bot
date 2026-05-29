import os
import re
import json
import tempfile
import requests
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Brand colors ──────────────────────────────────────────────────────────────
WINE       = RGBColor(0x7A, 0x00, 0x13)
GOLD       = RGBColor(0xD4, 0xA3, 0x47)
CHOCOLATE  = RGBColor(0x2B, 0x1B, 0x18)
CREAM      = RGBColor(0xFA, 0xF5, 0xEF)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Config ────────────────────────────────────────────────────────────────────
GROQ_API_KEY    = os.environ.get("GROQ_API_KEY", "")
TAVILY_API_KEY  = os.environ.get("TAVILY_API_KEY", "")
SLACK_BOT_TOKEN = os.environ.get("SLACK_BOT_TOKEN", "")

groq_client = Groq(api_key=GROQ_API_KEY) if GROQ_API_KEY else None

# Use 70b only for the client guide (quality-critical).
# Selling, digital product, and lead magnet use 8b-instant which has a
# SEPARATE 500k token/day limit — avoids burning the 100k/day 70b quota.
MODEL_MAIN = "llama-3.3-70b-versatile"
MODEL_FAST = "llama-3.1-8b-instant"

BRAND_DOCS_DIR = os.path.join(os.path.dirname(__file__), "brand_docs")

_GUIDE_VOICE_FILES = [
    "N2F CONTENT VOICE AND NAMING GUIDE_v1.md",
    "N2F Proprietary Coaching Mechanisms_v1.md",
    "Confident Yet Soft Brand Guide Mar 2026_v1.md",
    "markdown Premium Buyer OUTCOMES.md",
    "Confident Yet Soft PREMIUM Buyer MASTER SETUP PROMPT_v1.md",
]
_SELLING_FILES = [
    "N2F CONTENT VOICE AND NAMING GUIDE_v1.md",
    "High Ticket Buyer Content Guides_v1.md",
    "markdown Premium Buyer OUTCOMES.md",
    "ConfidentYetSoft_Premium_Buyers_PRISM_Content_Strategy_v1.md",
    "Leanne M Rich Buyer Guidance_v1.md",
]

# ── Brand doc loading ─────────────────────────────────────────────────────────
_BRAND_DOCS = {}  # type: dict

def _load_brand_docs():
    if not os.path.exists(BRAND_DOCS_DIR):
        print(f"brand_docs/ not found at {BRAND_DOCS_DIR}")
        return
    for fname in os.listdir(BRAND_DOCS_DIR):
        fpath = os.path.join(BRAND_DOCS_DIR, fname)
        if not os.path.isfile(fpath):
            continue
        try:
            if fname.endswith('.docx'):
                doc = Document(fpath)
                text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
                _BRAND_DOCS[fname] = text
            elif fname.endswith(('.md', '.txt')):
                with open(fpath, 'r', encoding='utf-8', errors='ignore') as f:
                    _BRAND_DOCS[fname] = f.read()
        except Exception as e:
            print(f"Error loading brand doc {fname}: {e}")
    print(f"Loaded {len(_BRAND_DOCS)} brand docs")

_load_brand_docs()


def _get_brand_context(file_list, max_chars_each=4000):
    parts = []
    for fname in file_list:
        text = _BRAND_DOCS.get(fname, "")
        if text:
            parts.append(f"=== {fname} ===\n{text[:max_chars_each]}")
    return '\n\n'.join(parts)


# ── Groq helpers ──────────────────────────────────────────────────────────────
def _groq(prompt, temperature=0.4, max_tokens=4096, model=None):
    if model is None:
        model = MODEL_MAIN
    completion = groq_client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=temperature,
        max_tokens=max_tokens,
    )
    return completion.choices[0].message.content


def _parse_json(raw):
    text = re.sub(r'^```(?:json)?\s*\n?', '', raw.strip(), flags=re.MULTILINE)
    text = re.sub(r'\n?```\s*$', '', text, flags=re.MULTILINE).strip()
    try:
        return json.loads(text)
    except Exception:
        m = re.search(r'\{[\s\S]*\}', text)
        if m:
            try:
                return json.loads(m.group(0))
            except Exception:
                pass
    return {}


# ── Task extraction ───────────────────────────────────────────────────────────
_GUIDE_NOUNS  = {'guide', 'plan', 'protocol', 'resource', 'worksheet', 'checklist',
                 'document', 'program', 'template', 'assessment', 'framework', 'roadmap'}
_CREATE_VERBS = {'create', 'send', 'prepare', 'write', 'build', 'make', 'develop',
                 'share', 'email', 'put together'}

def extract_guide_from_tasks(coach_tasks):
    for task in (coach_tasks or []):
        task_lower = task.lower()
        if not (any(v in task_lower for v in _CREATE_VERBS) and
                any(n in task_lower for n in _GUIDE_NOUNS)):
            continue
        m = re.search(
            r'(?:the|a|an)\s+([\w\s\-\'&,]+?'
            r'(?:Guide|Plan|Protocol|Resource|Worksheet|Checklist|Document|'
            r'Program|Template|Assessment|Framework|Roadmap))',
            task, re.IGNORECASE
        )
        guide_name = m.group(1).strip() if m else task.strip()
        return {"found": True, "guide_name": guide_name, "full_task": task}
    return {"found": False, "guide_name": "", "full_task": ""}


def detect_guide_promise(transcript, fathom_summary, client_name):
    combined = (fathom_summary or transcript or "")[:6000]
    prompt = f"""Analyze this call. Did the coach explicitly commit to creating and sending a guide, plan, or document to {client_name}?

CALL:
{combined}

Return ONLY valid JSON:
{{"promised": true_or_false, "topic": "exact topic (empty if false)", "client_context": "key call details (empty if false)"}}"""
    result = _parse_json(_groq(prompt, temperature=0.2, max_tokens=400, model=MODEL_FAST))
    if not isinstance(result, dict):
        return {"promised": False, "topic": "", "client_context": ""}
    result.setdefault("promised", False)
    result.setdefault("topic", "")
    result.setdefault("client_context", "")
    return result


# ── Market research ───────────────────────────────────────────────────────────
def run_guide_market_research(topic):
    if not TAVILY_API_KEY:
        print("TAVILY_API_KEY not set — skipping market research")
        return []
    queries = [
        f"PCOS {topic} questions women asking reddit quora 2025 2026",
        f"{topic} women hormonal health digital guide best selling 2025 2026",
        f"PCOS {topic} lead magnet trending instagram tiktok 2026",
    ]
    results = []
    for q in queries:
        try:
            resp = requests.post(
                "https://api.tavily.com/search",
                json={"api_key": TAVILY_API_KEY, "query": q, "search_depth": "basic",
                      "max_results": 4, "include_answer": True},
                timeout=15,
            )
            if resp.status_code == 200:
                data = resp.json()
                snippet = (data.get("answer", "") + "\n" +
                           "\n".join(f"- {r.get('title','')}: {r.get('content','')[:200]}"
                                     for r in data.get("results", [])[:3]))
                results.append({"query": q, "findings": snippet[:800]})
        except Exception as e:
            print(f"Tavily error '{q}': {e}")
    return results


# ── Step 3: Client guide ──────────────────────────────────────────────────────
def generate_client_guide(topic, client_context, client_name,
                          transcript="", fathom_summary=""):
    """
    Returns richly structured guide data for docx rendering.
    Uses MODEL_MAIN (70b) for quality. Sections are prescribed by heading to
    ensure consistent, high-touch structure across all cultural food guides.
    """
    import time as _time

    # Compact voice rules — avoids sending 17k chars of brand docs
    VOICE = (
        "Tone: Confident Yet Soft. Warm, direct, knowledgeable best friend — not a textbook. "
        "Outcome-first language: lead with transformation, not process. "
        "Client is high-ticket: WE do the research for her, we hand her cheat sheets. "
        "Never tell her to 'go research' anything. "
        "BANNED PHRASES: 'culturally intelligent', 'culturally aware', 'culturally sensitive', "
        "'hormone journey', 'health journey', 'empower yourself', 'take control of your health', "
        "'holistic', 'portion control', 'mindful eating', 'food scale', 'measuring cups'. "
        "N2F brand: New To Fitness by Elaine Acheampong (@ilaineyflex). "
        "CTAs: private WhatsApp community, booking a call, follow @ilaineyflex."
    )

    transcript_excerpt = (transcript or fathom_summary or "")[:4000]

    # Only reference mechanisms the client actually uses
    mechanism_check = ""
    for mech in ["Hormone-First Approach", "HormoneSync Movement",
                 "10-Hour Hormone Optimization Window", "Hormonal Driver Assessment"]:
        if mech.lower() in transcript_excerpt.lower():
            mechanism_check += f"\n- {mech}"
    if not mechanism_check:
        mechanism_check = "None explicitly mentioned in this call — do NOT introduce any N2F mechanism names."

    prompt = f"""You are writing a comprehensive, high-touch client-facing guide for New To Fitness.

VOICE: {VOICE}

TOPIC: "{topic}"
CLIENT: {client_name}
CALL NOTES: {client_context}

CALL TRANSCRIPT (extract every specific example, food swap, or tip Elaine gave and weave it into the guide):
{transcript_excerpt}

N2F MECHANISMS CLEARED (only reference if listed):
{mechanism_check}

SCIENTIFIC STANDARD: Back every recommendation with real science. Include 2-3 statistics or study references in the intro (e.g. "Women with PCOS have up to 4x higher insulin resistance..."). Action steps must cite the physiological mechanism behind the recommendation (e.g. "Switching to basmati from jasmine lowers the glycemic response by ~30% because of its higher amylose content").

HABIT COACHING STANDARD: Frame every action step using proven habit-change psychology — make it small, specific, tied to an existing habit, and immediately achievable. Reduce friction. Example: "Next time you're making jollof, swap out 1 cup of rice for 1 cup of coarse bulgur — keep everything else exactly the same."

SECTION STRUCTURE — generate ALL 9 sections below in this exact order. Each section must have a `section_intro` (2-3 sentences explaining what this section is and why it matters before diving into content), then its content:

SECTION 1 — "Why This Guide Exists: Cultural Food + PCOS"
Include: 2-3 PCOS/insulin resistance statistics, a relatable story hook (e.g. "You've eaten jollof your whole life and your body never complained — until now"), and what the client will be able to do differently after this guide. This is the emotional and scientific foundation.

SECTION 2 — "Rice & Your Blood Sugar: The Most Impactful Swap You Can Make"
Include: GI comparison table for Jasmine rice, Basmati rice, Sticky/Glutinous rice, Brown Jasmine, Brown Basmati, and Parboiled rice. Add a key disclaimer that GI is a tool but doesn't account for food pairing (protein, fat, fiber lower the actual blood sugar response). Then: a table of blood-sugar-friendly grain alternatives (Coarse Bulgur, Farro, Couscous, Quinoa, Barley, Pearl Millet) with columns: Grain | GI Score | Flavor Profile | Best Used In | Why It Works For PCOS. NOTE: For African dishes like jollof or stew-based dishes, coarse bulgur has the most similar texture to rice — it soaks up sauces the same way and maintains consistency. This is the coach-recommended #1 swap for African cultural dishes.

SECTION 3 — "The Hormone-Friendly Plate: Your Cultural Food Formula"
Include: the plate formula (1/2 plate colorful vegetables or fruit, 1/4 slow-digesting carbs WITH EXAMPLES, 1/4 to 1/2 plate protein WITH EXAMPLES, cook in healthy fat, water or sparkling water). Then generate EXACTLY 7 detailed cultural cuisine examples of a Hormone-Friendly Plate. For EACH example: name the cuisine, describe exactly what is on the plate (specific dishes), and write a Gemini AI image generation prompt for a realistic editorial photo of that plate of food. Cuisines to cover: West African, East African, South Asian (Indian/Pakistani), Caribbean, Middle Eastern, South American, Southern US/Soul Food.

SECTION 4 — "Cultural Food Ingredient Swaps: Your Coach-Vetted Cheat Sheet"
This is the research WE did so she doesn't have to. Generate a comprehensive swap table organized by cuisine region. NEVER say "research alternatives." The table must have 4 columns: Traditional Ingredient | PCOS-Friendly Alternative | Hormonal/Physiological Benefit | Outcome For You (client-facing benefit). Cover: West African ingredients (palm oil, Maggi/seasoning cubes, white rice, deep-fried items, full-fat coconut milk), South Asian (white rice, ghee in excess, refined flour/maida, full-fat yogurt in excess), East/Southeast Asian (soy sauce high sodium, white rice, deep-fried proteins), Middle Eastern (white rice, tahini in excess, refined pita), South American (white rice, refined corn products), Southern US/Soul Food (lard/crisco, white rice, deep-fried proteins, pork products).

SECTION 5 — "Making Cultural Carbs Work For You: Beyond Rice"
Focus on plantain (common in African, Caribbean, South American cuisines): ripe vs less ripe plantain and the blood sugar implications (ripe plantain = higher GI). Give cooking method swaps: deep fry → air fry (same crispy result), or boil in water and serve with stew and protein. Add examples for other cultural carbs: fufu/pounded yam alternatives, eba/garri modifications, roti/chapati swaps, injera modifications, tortilla alternatives.

SECTION 6 — "Protein in Cultural Foods: Why Your Plate Is Unbalanced and How to Fix It"
Open with the specific problem: many cultural meals default to 3/4 plate of rice/carbs, a side of plantain, and 1 small piece of protein — this pattern causes the blood sugar crashes, energy dips, and cravings women with PCOS experience. Explain the science of protein's role in blood sugar regulation. Then: how to increase protein portions in traditional dishes. Generate a table: Protein Type | Standard Cut | Hormone-Friendly Cut | Why (e.g. Extra Lean Ground Beef vs Regular Ground Beef, Chicken Breast vs Chicken Leg, Goat Leg vs Goat Shoulder, etc.). Include tips on cooking method: air fry vs deep fry for proteins. Include specific tip on high-fat meats: for meats like bacon where a low-sodium version isn't available, soak in cold water for 4-6 hours to draw out most of the salt, then air fry until crispy to render out the majority of the fat — this makes it significantly more manageable.

SECTION 7 — "Cooking Equipment Worth Investing In"
Equipment that supports the cooking method swaps recommended in this guide. Include: Air fryer (most important — why it works for cultural cooking), Non-stick cookware (explain PFOA-free options), Stainless steel cookware (when to use), Cast iron pan (benefits for cultural cooking, adds iron), Other relevant tools. For each: brief explanation of why it matters for PCOS-friendly cultural cooking.

SECTION 8 — "Your Cultural Food Swaps Grocery List"
A categorized grocery list of all the swap items referenced in this guide. Format as categories with items. Each item should note what it replaces. Categories: Grains & Carbs, Oils & Fats, Proteins, Vegetables & Produce, Seasonings & Condiments, Cooking Equipment.

SECTION 9 — "Your 30-Day Cultural Food Audit: A Step-by-Step Plan"
Frame using the habit psychology principle of "minimum viable change" — they do NOT need to throw out their pantry or replace everything at once. Step-by-step weekly plan: Week 1 (audit only: write down your 5 most-eaten cultural dishes and their main ingredients), Week 2 (pick ONE swap to implement — whichever comes up naturally in your next grocery shop), Week 3-4 (add a second swap, let the coach/group chat know what you're working on). End with: "At the end of 30 days, message your coach or post in the group chat: what is the ONE cultural food swap you have implemented?" and direct them to share their main takeaway from this guide with their coach.

ADDITIONAL OUTPUTS:
- `trainerize_habits`: List 6-8 specific, trackable habits the coach can add in Trainerize for any client using this guide. Format as short action habits (e.g. "Build a Hormone-Friendly Plate at dinner", "Swap rice for bulgur in 1 meal this week"). These should be implementable immediately.
- `quick_reference_checklist`: 10-12 key actions from the entire guide, phrased as checkboxes
- `note_to_send_to_client`: 3-4 sentence warm message from Elaine to {client_name} referencing something specific from their call. Confident Yet Soft.
- `keyword`: A specific 2-3 word phrase someone would type to trigger this guide automation (not: food, cook, eat, oil)
- `references`: List of 3-5 scientific sources cited (journal/study name + finding, no full URLs needed)

Return ONLY valid JSON with no markdown fences. Use \\n for line breaks inside strings, never literal newlines:
{{
  "guide_name": "Outcome-first guide name (N2F naming conventions)",
  "subtitle": "What they will achieve — specific and client-facing",
  "problem_statement": "One sentence: the exact problem this guide solves",
  "intro_science": "2-3 sentences with specific PCOS statistics and the science hook",
  "intro_story": "2-3 sentences: a relatable story hook the client has likely lived",
  "intro_promise": "1-2 sentences: what they will be able to do differently after this guide",
  "sections": [
    {{
      "heading": "Exact section heading as prescribed above",
      "section_intro": "2-3 sentences explaining what this section is and why it matters — before any content",
      "key_insight_box": "The single most important insight from this section, 1-2 sentences",
      "action_steps": ["Specific actionable step with exact method/amount/timing", "Another step"],
      "table": {{
        "use": false,
        "caption": "",
        "headers": [],
        "rows": []
      }},
      "cuisine_plate_examples": [],
      "key_takeaway": "One sentence: what to remember and implement",
      "image_prompt": "Gemini AI image generation prompt: realistic editorial wellness photo of [specific scene]",
      "review_note": "What Elaine should verify — empty string if confident"
    }}
  ],
  "quick_reference_checklist": ["Checkbox action 1", "Checkbox action 2"],
  "trainerize_habits": ["Trackable habit 1 for Trainerize", "Trackable habit 2"],
  "mid_doc_cta": "Warm mid-document CTA — private WhatsApp community or next call (1-2 sentences)",
  "end_cta": "Closing CTA — book a call + join community + follow @ilaineyflex (2-3 sentences)",
  "note_to_send_to_client": "3-4 sentences from Elaine to {client_name}, referencing something specific from their call",
  "references": ["Study/source 1", "Study/source 2", "Study/source 3"],
  "keyword": "specific trigger phrase"
}}

For Section 3, the `cuisine_plate_examples` field must be an array of objects:
{{"cuisine": "West African", "plate_description": "exactly what is on this plate", "image_prompt": "Gemini AI prompt for this specific plate"}}

For the GROCERY LIST section (Section 8), use `action_steps` as a flat list of items formatted as "Item name [replaces: traditional ingredient]". Use `table` for any category breakdown."""

    raw = _groq(prompt, temperature=0.4, max_tokens=8000, model=MODEL_MAIN)
    result = _parse_json(raw)
    if not result or not result.get("sections"):
        result = {
            "guide_name": f"{topic} Guide",
            "subtitle": "Hormone-friendly strategies for your cultural kitchen",
            "problem_statement": f"How to navigate {topic} with PCOS without giving up your food.",
            "intro_science": f"Women with PCOS have up to 4x higher insulin resistance than women without the condition.",
            "intro_story": "You've been eating the foods you grew up with your whole life — and now you're being told they're the problem.",
            "intro_promise": "This guide gives you the exact swaps, strategies, and cheat sheets to keep eating your cultural foods while supporting your hormones.",
            "sections": [{"heading": "Overview", "section_intro": "Content generation failed — regenerate.", "content": raw[:2000],
                          "key_insight_box": "", "action_steps": [],
                          "table": {"use": False, "caption": "", "headers": [], "rows": []},
                          "cuisine_plate_examples": [],
                          "key_takeaway": "", "image_prompt": "", "review_note": "Full review needed"}],
            "quick_reference_checklist": [],
            "trainerize_habits": [],
            "mid_doc_cta": "You don't have to figure this out alone — join our private WhatsApp community for ongoing support and accountability.",
            "end_cta": "Ready to go deeper? Book your next call with Elaine at newtofitness.com. Follow @ilaineyflex for daily tips.",
            "note_to_send_to_client": f"Hi {client_name}, here is the guide we discussed on our call.",
            "references": [],
            "keyword": "cultural food swaps",
        }
    return result


# ── Step 4: Selling opportunities ─────────────────────────────────────────────
def generate_selling_opportunities_content(guide_data, topic, search_results):
    brand_ctx = _get_brand_context(_SELLING_FILES, max_chars_each=3500)
    guide_name = guide_data.get("guide_name", topic)
    search_text = ("\n".join(f"[Search: {r['query']}]\n{r['findings']}"
                             for r in search_results)
                   if search_results else "No live search data.")

    prompt = f"""You are a digital product strategist for New To Fitness (N2F), Elaine Acheampong's premium PCOS coaching brand.

BRAND GUIDANCE:
{brand_ctx}

GUIDE CREATED:
Name: {guide_name}
Topic: {topic}

MARKET RESEARCH:
{search_text}

CRITICAL CONSTRAINTS:
- Digital products are LOW-TICKET DOWNLOADABLE GUIDES: $10-$47 price range ONLY
- Each product solves ONE SPECIFIC PROBLEM and delivers a quick win
- They are NOT coaching programs, NOT courses, NOT broad lifestyle guides
- Think: a beautiful 10-20 page PDF guide someone downloads, reads in 30 min, and implements today
- The guide must be complete enough to deliver a real result on its own

PRODUCT NAMING RULES (from N2F naming guide):
- Lead with transformation, not topic
- Premium buyer language — outcome-first
- No clinical jargon, no generic wellness phrases
- Title: what you do. Subtitle: what you get.

LEAD MAGNET RULES:
- Free, shorter version that delivers a quick win but creates desire for the paid guide
- Does NOT share a name with any top-5 upsell product (never give away for free what you charge for)
- Designed to pull people into our private WhatsApp community or social media

Return ONLY valid JSON with no markdown fences:
{{
  "upsell_products": [
    {{
      "rank": 1,
      "name": "Transformation-first product name",
      "subheader": "Outcome-focused subheader",
      "price_point": "$X (must be $10-$47)",
      "pages_estimate": "10-20 pages",
      "quick_win": "What result they get within 1-2 weeks",
      "score": 95,
      "reasoning": "Why this will sell based on market research"
    }}
  ],
  "lead_magnets": [
    {{
      "rank": 1,
      "name": "Lead magnet name",
      "subheader": "Benefit subheader",
      "format": "checklist/mini-guide/swap list/etc",
      "platform": "Instagram/TikTok/Both",
      "engagement_prediction": "Why this performs well on that platform"
    }}
  ],
  "carousel_scripts": [
    {{
      "promotes": "product or lead magnet name",
      "caption": "Full Instagram caption: hook, body, CTA. Include hashtags.",
      "slides": ["Slide 1: Hook copy (5-10 words)", "Slide 2 copy", "Slide 3 copy", "Slide 4 copy", "Slide 5 copy", "Slide 6: CTA copy"]
    }}
  ],
  "automation_keyword": "Specific 2-3 word phrase unique to this guide topic"
}}

Include 3-5 upsell products, 3-5 lead magnets, 5 carousel scripts."""

    raw = _groq(prompt, temperature=0.5, max_tokens=5000, model=MODEL_FAST)
    result = _parse_json(raw)
    if not isinstance(result, dict):
        result = {"upsell_products": [], "lead_magnets": [],
                  "carousel_scripts": [], "automation_keyword": ""}
    return result


# ── Step 5: Digital product ───────────────────────────────────────────────────
def generate_digital_product_content(guide_data, selling_data,
                                     transcript="", fathom_summary=""):
    brand_ctx = _get_brand_context(_GUIDE_VOICE_FILES, max_chars_each=3000)
    upsells = selling_data.get("upsell_products", [])
    top = upsells[0] if upsells else {}
    product_name = top.get("name", guide_data.get("guide_name", ""))
    subheader    = top.get("subheader", "")
    price_point  = top.get("price_point", "$27")
    quick_win    = top.get("quick_win", "")
    topic        = guide_data.get("guide_name", "")
    transcript_excerpt = (transcript or fathom_summary or "")[:3000]

    prompt = f"""You are creating a LOW-TICKET DIGITAL PRODUCT for New To Fitness by Elaine Acheampong.

BRAND VOICE:
{brand_ctx}

PRODUCT:
Name: {product_name}
Subheader: {subheader}
Price: {price_point}
Quick win: {quick_win}
Topic: {topic}

CALL CONTENT TO DRAW FROM (use specific examples from this call as the foundation):
{transcript_excerpt}

CRITICAL RULES:
- This is a downloadable PDF guide, $10-47, 10-20 pages
- ONE specific problem, ONE clear transformation, ONE quick win
- NOT a coaching program. NOT a course. NOT broad lifestyle advice.
- Every section needs SPECIFIC ACTIONABLE STEPS with exact methods/ratios/timing
- Significantly more detailed than the free client guide but still focused on the same ONE topic
- Include at least one swap table or comparison table
- Include a simple 7-day implementation checklist (not a complex plan)
- Sections should have: key insight callout, action steps, table where relevant

BANNED: "empower", "journey", "holistic", "culturally intelligent", broad nutrition advice unrelated to the specific topic

Return ONLY valid JSON with no markdown fences:
{{
  "product_name": "{product_name}",
  "subheader": "{subheader}",
  "price_point": "{price_point}",
  "quick_win": "{quick_win}",
  "social_proof_placeholder": "Description of the type of before/after story that would work here (e.g. 'Client who kept eating jollof rice and still lost X lbs in Y weeks')",
  "intro": "2-3 sentence outcome-first welcome. Warm and direct.",
  "sections": [
    {{
      "heading": "Section heading",
      "key_insight_box": "The most important insight from this section in 1-2 sentences",
      "content": "2-3 paragraphs of specific, actionable content",
      "action_steps": ["Specific step 1", "Specific step 2"],
      "table": {{
        "use": false,
        "caption": "",
        "headers": [],
        "rows": []
      }},
      "image_prompt": "Gemini AI image prompt: realistic editorial photo of [specific scene]",
      "review_note": "What Elaine should verify — empty if confident"
    }}
  ],
  "seven_day_checklist": [
    {{"day": "Day 1-2", "action": "Specific action to take"}},
    {{"day": "Day 3-4", "action": "Next action"}},
    {{"day": "Day 5-7", "action": "Final action"}}
  ],
  "mid_cta": "Mid-document CTA: invite to private WhatsApp community (2 sentences)",
  "end_cta": "Closing CTA: book a discovery call + follow @ilaineyflex + join community (3 sentences, warm and direct)"
}}"""

    raw = _groq(prompt, temperature=0.4, max_tokens=6000, model=MODEL_FAST)
    result = _parse_json(raw)
    if not isinstance(result, dict) or not result.get("sections"):
        result = {
            "product_name": product_name, "subheader": subheader,
            "price_point": price_point, "quick_win": quick_win,
            "social_proof_placeholder": "Client before/after story",
            "intro": f"Welcome to {product_name}.",
            "sections": [{"heading": "Content", "content": raw[:3000],
                          "key_insight_box": "", "action_steps": [],
                          "table": {"use": False, "caption": "", "headers": [], "rows": []},
                          "image_prompt": "", "review_note": "Full review needed"}],
            "seven_day_checklist": [],
            "mid_cta": "Join our private WhatsApp community.",
            "end_cta": "Book your call at newtofitness.com."
        }
    return result


# ── Step 6: Lead magnet ───────────────────────────────────────────────────────
def generate_lead_magnet_content(guide_data, selling_data,
                                 transcript="", fathom_summary=""):
    brand_ctx = _get_brand_context(_GUIDE_VOICE_FILES, max_chars_each=2500)
    lead_magnets = selling_data.get("lead_magnets", [])
    top = lead_magnets[0] if lead_magnets else {}
    freebie_name = top.get("name", f"{guide_data.get('guide_name', 'Guide')} — Free Guide")
    subheader    = top.get("subheader", "")
    fmt          = top.get("format", "mini-guide")
    topic        = guide_data.get("guide_name", "")
    transcript_excerpt = (transcript or fathom_summary or "")[:2000]

    prompt = f"""You are creating a FREE LEAD MAGNET for New To Fitness by Elaine Acheampong.

BRAND VOICE:
{brand_ctx}

LEAD MAGNET:
Name: {freebie_name}
Subheader: {subheader}
Format: {fmt}
Topic: {topic}

CALL CONTENT TO DRAW FROM:
{transcript_excerpt}

RULES:
- Genuinely valuable: delivers one quick win the reader can implement TODAY
- Deliberately incomplete: creates desire for the full paid guide
- 3-4 sections maximum, 5-8 pages
- Must include at least one swap list or comparison table
- Every section has specific actionable steps, not theory
- Ends with clear CTAs: join private WhatsApp community + follow @ilaineyflex + tease the paid guide
- Banned phrases: same as above

Return ONLY valid JSON with no markdown fences:
{{
  "freebie_name": "{freebie_name}",
  "subheader": "{subheader}",
  "intro": "1-2 sentences: what they get and why it matters right now",
  "sections": [
    {{
      "heading": "Section heading",
      "key_insight_box": "Key insight to highlight",
      "action_steps": ["Specific step 1", "Specific step 2"],
      "table": {{
        "use": false,
        "caption": "",
        "headers": [],
        "rows": []
      }},
      "image_prompt": "Gemini AI image prompt for this section",
      "review_note": ""
    }}
  ],
  "quick_checklist": ["Check item 1", "Check item 2"],
  "cta_community": "Invite to join N2F private WhatsApp community (1-2 sentences, warm)",
  "cta_upgrade": "Tease the paid guide — what more they get and how to get it (1-2 sentences)",
  "cta_social": "Invite to follow @ilaineyflex for daily tips (1 sentence)"
}}"""

    raw = _groq(prompt, temperature=0.4, max_tokens=3500, model=MODEL_FAST)
    result = _parse_json(raw)
    if not isinstance(result, dict) or not result.get("sections"):
        result = {
            "freebie_name": freebie_name, "subheader": subheader,
            "intro": f"Your free guide to {topic}.",
            "sections": [{"heading": "Content", "content": raw[:2000],
                          "key_insight_box": "", "action_steps": [],
                          "table": {"use": False, "caption": "", "headers": [], "rows": []},
                          "image_prompt": "", "review_note": "Full review needed"}],
            "quick_checklist": [],
            "cta_community": "Join our private WhatsApp community.",
            "cta_upgrade": "Ready for the full guide?",
            "cta_social": "Follow @ilaineyflex for daily tips."
        }
    return result


# ── Word doc helpers ───────────────────────────────────────────────────────────
def _set_shading(para, fill_hex):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.lstrip('#'))
    pPr.append(shd)


def _set_cell_shading(cell, fill_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex.lstrip('#'))
    tcPr.append(shd)


def _set_para_border(para, side, color_hex, size='12', space='4'):
    pPr = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn('w:pBdr'))
    if pBdr is None:
        pBdr = OxmlElement('w:pBdr')
        pPr.append(pBdr)
    el = OxmlElement(f'w:{side}')
    el.set(qn('w:val'), 'single')
    el.set(qn('w:sz'), size)
    el.set(qn('w:space'), space)
    el.set(qn('w:color'), color_hex.lstrip('#'))
    pBdr.append(el)


def _gold_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)
    _set_para_border(p, 'bottom', 'D4A347', size='8')


def _setup_doc(doc):
    style = doc.styles['Normal']
    style.font.name  = 'Calibri'
    style.font.size  = Pt(11)
    style.font.color.rgb = CHOCOLATE
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)


def _tmppath(prefix):
    fd, path = tempfile.mkstemp(suffix='.docx', prefix=prefix)
    os.close(fd)
    return path


# ── Reusable visual blocks ────────────────────────────────────────────────────

def _cover_page(doc, title, subtitle, tagline="New To Fitness | @ilaineyflex"):
    """Wine banner cover page."""
    # Top brand strip
    top = doc.add_paragraph()
    _set_shading(top, '7A0013')
    r = top.add_run(f"  {tagline}  ")
    r.font.size  = Pt(10)
    r.font.bold  = True
    r.font.color.rgb = GOLD
    top.paragraph_format.space_before = Pt(0)
    top.paragraph_format.space_after  = Pt(0)

    doc.add_paragraph()

    # Image placeholder
    img_box = doc.add_paragraph()
    _set_shading(img_box, 'F0E8D8')
    r = img_box.add_run(
        "\n[COVER IMAGE PLACEHOLDER]\n"
        "Gemini AI Prompt: \"Realistic editorial wellness photo of a woman preparing "
        "a vibrant African or South Asian meal in a modern kitchen, warm natural "
        "lighting, Vogue wellness aesthetic, N2F brand colors wine and gold accents\"\n"
        "Suggested size: full width, 3 inches tall\n"
    )
    r.font.size      = Pt(9)
    r.font.italic    = True
    r.font.color.rgb = RGBColor(0x7A, 0x6A, 0x5A)
    img_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_box.paragraph_format.space_after = Pt(12)

    # Title
    title_p = doc.add_paragraph()
    tr = title_p.add_run(title)
    tr.font.size  = Pt(26)
    tr.font.bold  = True
    tr.font.color.rgb = WINE
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(6)

    # Subtitle
    sub_p = doc.add_paragraph()
    sr = sub_p.add_run(subtitle)
    sr.font.size      = Pt(13)
    sr.font.italic    = True
    sr.font.color.rgb = CHOCOLATE
    sub_p.alignment   = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(20)

    _gold_rule(doc)
    doc.add_page_break()


def _section_heading(doc, heading):
    h = doc.add_paragraph()
    r = h.add_run(heading)
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = WINE
    h.paragraph_format.space_before = Pt(16)
    h.paragraph_format.space_after  = Pt(4)
    _set_para_border(h, 'bottom', 'D4A347', size='4')


def _key_insight_box(doc, text):
    """Gold left-border insight callout."""
    if not text:
        return
    p = doc.add_paragraph()
    _set_shading(p, 'FBF5E6')
    _set_para_border(p, 'left', 'D4A347', size='24', space='8')
    r = p.add_run(f"Key Insight: {text}")
    r.font.size      = Pt(10.5)
    r.font.italic    = True
    r.font.color.rgb = CHOCOLATE
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(8)


def _action_steps(doc, steps):
    """Formatted action step list with wine bullet markers."""
    if not steps:
        return
    label = doc.add_paragraph()
    lr = label.add_run("Action Steps")
    lr.font.size  = Pt(10)
    lr.font.bold  = True
    lr.font.color.rgb = WINE
    label.paragraph_format.space_before = Pt(8)
    label.paragraph_format.space_after  = Pt(2)

    for step in steps:
        p = doc.add_paragraph()
        r = p.add_run(f"•  {step}")
        r.font.size      = Pt(10.5)
        r.font.color.rgb = CHOCOLATE
        p.paragraph_format.left_indent   = Inches(0.2)
        p.paragraph_format.space_before  = Pt(2)
        p.paragraph_format.space_after   = Pt(3)


def _key_takeaway(doc, text):
    """Wine-background key takeaway box."""
    if not text:
        return
    p = doc.add_paragraph()
    _set_shading(p, 'F7E8E8')
    r = p.add_run(f"✨  Remember: {text}")
    r.font.size      = Pt(10.5)
    r.font.bold      = True
    r.font.color.rgb = WINE
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(10)


def _add_table(doc, headers, rows, caption=""):
    """Wine-header table."""
    if not headers and not rows:
        return
    n_cols  = max(len(headers), max((len(r) for r in rows), default=0)) if headers or rows else 2
    n_rows  = len(rows) + (1 if headers else 0)
    if n_cols == 0 or n_rows == 0:
        return

    if caption:
        cap = doc.add_paragraph()
        cr  = cap.add_run(caption)
        cr.font.size      = Pt(9)
        cr.font.italic    = True
        cr.font.color.rgb = GOLD
        cap.paragraph_format.space_before = Pt(8)
        cap.paragraph_format.space_after  = Pt(2)

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.style = 'Table Grid'
    row_i = 0

    if headers:
        for i, h_text in enumerate(headers[:n_cols]):
            cell = table.rows[0].cells[i]
            cell.text = h_text
            _set_cell_shading(cell, '7A0013')
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold      = True
                    run.font.color.rgb = WHITE
                    run.font.size      = Pt(10)
        row_i = 1

    for row_data in rows:
        for i, val in enumerate(row_data[:n_cols]):
            cell = table.rows[row_i].cells[i]
            cell.text = str(val)
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size      = Pt(10)
                    run.font.color.rgb = CHOCOLATE
        row_i += 1

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def _image_placeholder(doc, gemini_prompt):
    """Gray placeholder box with Gemini image prompt."""
    if not gemini_prompt:
        return
    p = doc.add_paragraph()
    _set_shading(p, 'ECECEC')
    r = p.add_run(
        f"[IMAGE PLACEHOLDER]\n"
        f"Gemini AI Prompt: {gemini_prompt}\n"
        f"Suggested size: full width, 2.5 inches tall"
    )
    r.font.size      = Pt(8.5)
    r.font.italic    = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(10)


def _cta_block(doc, text, style='mid'):
    """
    style='mid': gold background subtle CTA
    style='end': wine background bold closing CTA
    """
    doc.add_paragraph()
    p = doc.add_paragraph()
    fill = 'FBF5E6' if style == 'mid' else '7A0013'
    _set_shading(p, fill)
    _set_para_border(p, 'left', 'D4A347' if style == 'mid' else 'D4A347', size='20', space='8')
    r = p.add_run(text)
    r.font.size      = Pt(10.5)
    r.font.bold      = (style == 'end')
    r.font.color.rgb = CHOCOLATE if style == 'mid' else WHITE
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)


def _checklist(doc, items, heading="Quick Reference Checklist"):
    if not items:
        return
    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run(heading)
    hr.font.size  = Pt(13)
    hr.font.bold  = True
    hr.font.color.rgb = WINE
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(6)
    _set_para_border(h, 'bottom', 'D4A347', size='4')

    for item in items:
        p = doc.add_paragraph()
        _set_shading(p, 'FAFAF5')
        r = p.add_run(f"☐  {item}")
        r.font.size      = Pt(10.5)
        r.font.color.rgb = CHOCOLATE
        p.paragraph_format.left_indent   = Inches(0.2)
        p.paragraph_format.space_before  = Pt(3)
        p.paragraph_format.space_after   = Pt(3)


def _social_proof_placeholder(doc, description):
    if not description:
        return
    doc.add_paragraph()
    h = doc.add_paragraph()
    hr = h.add_run("Client Success Story")
    hr.font.size  = Pt(12)
    hr.font.bold  = True
    hr.font.color.rgb = WINE

    p = doc.add_paragraph()
    _set_shading(p, 'F0F7EE')
    _set_para_border(p, 'left', '4A7A4A', size='20', space='8')
    r = p.add_run(
        f"[SOCIAL PROOF PLACEHOLDER]\n"
        f"{description}\n\n"
        f"Replace with a real client before/after story, screenshot, or testimonial.\n"
        f"Format: Photo (if available) + 2-3 sentence story + measurable result."
    )
    r.font.size      = Pt(9)
    r.font.italic    = True
    r.font.color.rgb = RGBColor(0x2A, 0x5A, 0x2A)
    p.paragraph_format.left_indent  = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(12)


def _review_callout(doc, review_note):
    if not review_note:
        return
    p = doc.add_paragraph()
    _set_shading(p, 'FFF3CD')
    r = p.add_run(f"⚠ REVIEW FOR ELAINE: {review_note}")
    r.font.italic    = True
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x7A, 0x00, 0x13)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)


def _claude_voice_prompt(doc):
    p = doc.add_paragraph()
    _set_shading(p, 'EDF4FB')
    r = p.add_run(
        "\U0001f4cb TO REFINE BRAND VOICE: Copy any section into Claude with: "
        "\"Rewrite for N2F Confident Yet Soft premium buyer voice. "
        "Ambitious woman with PCOS, wants outcomes not theory. "
        "Outcome-first language, no jargon, warm but authoritative. "
        "Reference N2F mechanisms only if part of her current program.\""
    )
    r.font.italic    = True
    r.font.size      = Pt(9)
    r.font.color.rgb = RGBColor(0x1E, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)


def _footer_line(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size      = Pt(8)
    r.font.color.rgb = GOLD
    r.font.italic    = True
    p.alignment      = WD_ALIGN_PARAGRAPH.CENTER


# ── Docx builders ─────────────────────────────────────────────────────────────

def build_client_guide_docx(guide_data, client_name):
    doc = Document()
    _setup_doc(doc)

    # Cover
    _cover_page(doc,
                guide_data.get("guide_name", "Your Guide"),
                guide_data.get("subtitle", ""),
                "New To Fitness | @ilaineyflex")

    # Problem statement
    ps = guide_data.get("problem_statement", "")
    if ps:
        p = doc.add_paragraph()
        _set_shading(p, 'F7E8E8')
        r = p.add_run(ps)
        r.font.size      = Pt(11)
        r.font.bold      = True
        r.font.color.rgb = WINE
        p.paragraph_format.space_after = Pt(8)

    # Structured intro (science stat + story + promise)
    for intro_key in ("intro_science", "intro_story", "intro_promise"):
        text = guide_data.get(intro_key, "")
        if text:
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

    # Fallback for old-format single intro string
    intro = guide_data.get("intro", "")
    if intro and not guide_data.get("intro_science"):
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(10)

    # Sections
    for i, s in enumerate(guide_data.get("sections", []), 1):
        _section_heading(doc, s.get("heading", ""))

        # Section intro — explains what this section covers before content
        s_intro = s.get("section_intro", "")
        if s_intro:
            p = doc.add_paragraph(s_intro)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(8)

        _key_insight_box(doc, s.get("key_insight_box", ""))
        _action_steps(doc, s.get("action_steps", []))

        tbl = s.get("table", {})
        if tbl.get("use") and tbl.get("headers"):
            _add_table(doc, tbl["headers"], tbl.get("rows", []),
                       caption=tbl.get("caption", ""))

        # Cuisine plate examples (Section 3 — Hormone-Friendly Plate)
        plates = s.get("cuisine_plate_examples", [])
        if plates:
            ph = doc.add_paragraph()
            r = ph.add_run("Hormone-Friendly Plate Examples Across Cultural Cuisines")
            r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = WINE
            ph.paragraph_format.space_before = Pt(10)
            ph.paragraph_format.space_after  = Pt(4)
            for plate in plates:
                pc = doc.add_paragraph()
                rc = pc.add_run(plate.get("cuisine", "") + ": ")
                rc.font.bold = True; rc.font.color.rgb = WINE; rc.font.size = Pt(10.5)
                rt = pc.add_run(plate.get("plate_description", ""))
                rt.font.size = Pt(10.5); rt.font.color.rgb = CHOCOLATE
                pc.paragraph_format.space_after = Pt(3)
                _image_placeholder(doc, plate.get("image_prompt", ""))

        _key_takeaway(doc, s.get("key_takeaway", ""))
        _image_placeholder(doc, s.get("image_prompt", ""))
        _review_callout(doc, s.get("review_note", ""))

        # Mid-doc CTA after section 4
        if i == 4:
            _cta_block(doc, guide_data.get("mid_doc_cta", ""), style='mid')

    # Quick reference checklist
    _checklist(doc, guide_data.get("quick_reference_checklist", []))

    # References
    refs = guide_data.get("references", [])
    if refs:
        doc.add_paragraph()
        rh = doc.add_paragraph()
        r = rh.add_run("References")
        r.font.size = Pt(11); r.font.bold = True; r.font.color.rgb = WINE
        rh.paragraph_format.space_before = Pt(12)
        _set_para_border(rh, 'bottom', 'D4A347', size='4')
        for ref in refs:
            rp = doc.add_paragraph()
            rr = rp.add_run(f"• {ref}")
            rr.font.size = Pt(9); rr.font.italic = True; rr.font.color.rgb = CHOCOLATE
            rp.paragraph_format.left_indent = Inches(0.2)
            rp.paragraph_format.space_after = Pt(2)

    # Closing CTA
    _cta_block(doc, guide_data.get("end_cta", ""), style='end')

    _claude_voice_prompt(doc)
    doc.add_paragraph()
    _footer_line(doc, "New To Fitness | @ilaineyflex | Educational purposes only — not medical advice.")

    path = _tmppath('guide_')
    doc.save(path)
    return path


def build_selling_opportunities_docx(selling_data, guide_name):
    doc = Document()
    _setup_doc(doc)

    _cover_page(doc, f"{guide_name}", "Selling Opportunities Assessment",
                "New To Fitness — Internal Strategy Document")

    # Upsell products
    h = doc.add_paragraph()
    r = h.add_run("DIGITAL PRODUCT UPSELL OPPORTUNITIES")
    r.font.bold = True; r.font.size = Pt(14); r.font.color.rgb = WINE
    _gold_rule(doc)

    for prod in selling_data.get("upsell_products", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        r = p.add_run(f"#{prod.get('rank')}  {prod.get('name','')}")
        r.font.bold = True; r.font.size = Pt(13); r.font.color.rgb = WINE

        sub = doc.add_paragraph()
        sub.add_run(prod.get("subheader","")).font.italic = True

        meta_items = [
            f"Price: {prod.get('price_point','')}",
            f"Pages: {prod.get('pages_estimate','')}",
            f"Score: {prod.get('score','')}",
        ]
        meta = doc.add_paragraph()
        r = meta.add_run("  |  ".join(i for i in meta_items if i.split(': ')[1]))
        r.font.size = Pt(9.5); r.font.color.rgb = GOLD

        if prod.get("quick_win"):
            qw = doc.add_paragraph()
            _set_shading(qw, 'F0F7EE')
            r = qw.add_run(f"Quick win: {prod['quick_win']}")
            r.font.size = Pt(10); r.font.color.rgb = RGBColor(0x1A, 0x4A, 0x1A)

        reason = doc.add_paragraph(prod.get("reasoning",""))
        reason.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # Lead magnets
    h2 = doc.add_paragraph()
    r2 = h2.add_run("LEAD MAGNETS & FREEBIES")
    r2.font.bold = True; r2.font.size = Pt(14); r2.font.color.rgb = WINE
    _gold_rule(doc)

    for lm in selling_data.get("lead_magnets", []):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        r = p.add_run(f"#{lm.get('rank')}  {lm.get('name','')}")
        r.font.bold = True; r.font.size = Pt(12); r.font.color.rgb = CHOCOLATE

        sub = doc.add_paragraph()
        sub.add_run(lm.get("subheader","")).font.italic = True

        meta = doc.add_paragraph()
        r = meta.add_run(f"Format: {lm.get('format','')}  |  Platform: {lm.get('platform','')}")
        r.font.size = Pt(9.5); r.font.color.rgb = GOLD

        pred = doc.add_paragraph(lm.get("engagement_prediction",""))
        pred.paragraph_format.space_after = Pt(10)

    doc.add_paragraph()

    # Carousel scripts
    h3 = doc.add_paragraph()
    r3 = h3.add_run("CAROUSEL SCRIPTS (ready to use)")
    r3.font.bold = True; r3.font.size = Pt(14); r3.font.color.rgb = WINE
    _gold_rule(doc)

    for i, script in enumerate(selling_data.get("carousel_scripts", []), 1):
        cs = doc.add_paragraph()
        cs.paragraph_format.space_before = Pt(16)
        r = cs.add_run(f"CAROUSEL {i} — Promoting: {script.get('promotes','')}")
        r.font.bold = True; r.font.size = Pt(11)

        cap_lbl = doc.add_paragraph()
        r = cap_lbl.add_run("CAPTION:")
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD

        cap_p = doc.add_paragraph()
        _set_shading(cap_p, 'FAFAF0')
        cap_p.add_run(script.get("caption","")).font.size = Pt(10)
        cap_p.paragraph_format.space_after = Pt(6)

        slides_lbl = doc.add_paragraph()
        r = slides_lbl.add_run("SLIDE COPY:")
        r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = GOLD

        slides = script.get("slides", [])
        if slides:
            _add_table(doc,
                       ["Slide", "Copy"],
                       [[str(j+1), text] for j, text in enumerate(slides)])

    # Automation keyword
    doc.add_paragraph()
    kw = doc.add_paragraph()
    _set_shading(kw, 'F7E8E8')
    r = kw.add_run(f"AUTOMATION KEYWORD: {selling_data.get('automation_keyword','')}")
    r.font.bold = True; r.font.color.rgb = WINE

    path = _tmppath('selling_')
    doc.save(path)
    return path


def build_digital_product_docx(product_data):
    doc = Document()
    _setup_doc(doc)

    _cover_page(doc,
                product_data.get("product_name", "Guide"),
                product_data.get("subheader", ""),
                f"New To Fitness | @ilaineyflex  |  {product_data.get('price_point','')}")

    # Quick win promise
    qw = product_data.get("quick_win", "")
    if qw:
        p = doc.add_paragraph()
        _set_shading(p, 'F0F7EE')
        _set_para_border(p, 'left', '4A7A4A', size='20', space='8')
        r = p.add_run(f"Your Quick Win: {qw}")
        r.font.size = Pt(11); r.font.bold = True
        r.font.color.rgb = RGBColor(0x1A, 0x4A, 0x1A)
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_after  = Pt(10)

    # Social proof placeholder
    _social_proof_placeholder(doc, product_data.get("social_proof_placeholder", ""))

    # Intro
    intro = product_data.get("intro", "")
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(10)

    # Sections
    for i, s in enumerate(product_data.get("sections", []), 1):
        _section_heading(doc, s.get("heading", ""))
        _key_insight_box(doc, s.get("key_insight_box", ""))

        content = s.get("content", "")
        for line in content.split('\n'):
            if line.strip():
                p = doc.add_paragraph(line.strip())
                p.paragraph_format.space_after = Pt(5)

        _action_steps(doc, s.get("action_steps", []))

        tbl = s.get("table", {})
        if tbl.get("use") and tbl.get("headers"):
            _add_table(doc, tbl["headers"], tbl.get("rows", []),
                       caption=tbl.get("caption", ""))

        _image_placeholder(doc, s.get("image_prompt", ""))
        _review_callout(doc, s.get("review_note", ""))

        if i == 2:
            _cta_block(doc, product_data.get("mid_cta", ""), style='mid')

    # 7-day checklist
    checklist = product_data.get("seven_day_checklist", [])
    if checklist:
        doc.add_paragraph()
        ch = doc.add_paragraph()
        r = ch.add_run("Your 7-Day Implementation Plan")
        r.font.size = Pt(13); r.font.bold = True; r.font.color.rgb = WINE
        ch.paragraph_format.space_before = Pt(12)
        _set_para_border(ch, 'bottom', 'D4A347', size='4')

        _add_table(doc,
                   ["Timeframe", "Your Action"],
                   [[item.get("day",""), item.get("action","")] for item in checklist])

    # Closing CTA
    _cta_block(doc, product_data.get("end_cta", ""), style='end')

    _claude_voice_prompt(doc)
    doc.add_paragraph()
    _footer_line(doc, "New To Fitness | @ilaineyflex | Educational purposes only — not medical advice.")

    path = _tmppath('digital_product_')
    doc.save(path)
    return path


def build_lead_magnet_docx(freebie_data):
    doc = Document()
    _setup_doc(doc)

    _cover_page(doc,
                freebie_data.get("freebie_name", "Free Guide"),
                freebie_data.get("subheader", ""),
                "New To Fitness | @ilaineyflex  |  FREE")

    intro = freebie_data.get("intro", "")
    if intro:
        p = doc.add_paragraph(intro)
        p.paragraph_format.space_after = Pt(10)

    for i, s in enumerate(freebie_data.get("sections", []), 1):
        _section_heading(doc, s.get("heading", ""))
        _key_insight_box(doc, s.get("key_insight_box", ""))
        _action_steps(doc, s.get("action_steps", []))

        tbl = s.get("table", {})
        if tbl.get("use") and tbl.get("headers"):
            _add_table(doc, tbl["headers"], tbl.get("rows", []),
                       caption=tbl.get("caption", ""))

        _image_placeholder(doc, s.get("image_prompt", ""))
        _review_callout(doc, s.get("review_note", ""))

    # Quick checklist
    _checklist(doc, freebie_data.get("quick_checklist", []),
               heading="Your Quick Action Checklist")

    # CTA section
    doc.add_paragraph()
    cta_h = doc.add_paragraph()
    r = cta_h.add_run("Ready for More?")
    r.font.size = Pt(14); r.font.bold = True; r.font.color.rgb = WINE
    _gold_rule(doc)

    comm = freebie_data.get("cta_community", "")
    upg  = freebie_data.get("cta_upgrade", "")
    soc  = freebie_data.get("cta_social", "")

    if comm:
        _cta_block(doc, f"\U0001f4ac  {comm}", style='mid')
    if upg:
        _cta_block(doc, f"⭐  {upg}", style='end')
    if soc:
        p = doc.add_paragraph()
        r = p.add_run(f"\U0001f4f7  {soc}")
        r.font.size = Pt(10.5); r.font.italic = True; r.font.color.rgb = CHOCOLATE

    _claude_voice_prompt(doc)
    doc.add_paragraph()
    _footer_line(doc, "New To Fitness | @ilaineyflex | Free to share — not medical advice.")

    path = _tmppath('lead_magnet_')
    doc.save(path)
    return path


# ── Slack file upload ──────────────────────────────────────────────────────────
def upload_file_to_slack(channel_id, thread_ts, file_path, filename, comment=""):
    headers = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
    file_size = os.path.getsize(file_path)

    resp = requests.get(
        "https://slack.com/api/files.getUploadURLExternal",
        headers=headers,
        params={"filename": filename, "length": file_size},
    )
    data = resp.json()
    if not data.get("ok"):
        print(f"Slack getUploadURL failed: {data.get('error')}")
        return False

    with open(file_path, 'rb') as f:
        up = requests.post(data["upload_url"], data=f)
    if up.status_code not in (200, 204):
        print(f"Slack file upload failed: {up.status_code}")
        return False

    payload = {"files": [{"id": data["file_id"], "title": filename}],
               "channel_id": channel_id}
    if thread_ts:
        payload["thread_ts"] = thread_ts
    if comment:
        payload["initial_comment"] = comment

    complete = requests.post(
        "https://slack.com/api/files.completeUploadExternal",
        headers={**headers, "Content-Type": "application/json"},
        json=payload,
    )
    result = complete.json()
    print(f"Slack upload: ok={result.get('ok')} {result.get('error','')}")
    return result.get("ok", False)


# ── Full pipeline ──────────────────────────────────────────────────────────────
def run_guide_pipeline(client_name, topic, client_context, transcript, fathom_summary,
                       channel_id, thread_ts, projects_channel_id):
    print(f"\n=== Guide pipeline: {client_name} | {topic} ===")
    guide_path = selling_path = digital_path = lm_path = None
    try:
        # 1. Client guide
        print("Step 1: client guide...")
        guide_data  = generate_client_guide(topic, client_context, client_name,
                                            transcript, fathom_summary)
        guide_name  = guide_data.get("guide_name", topic)
        note        = guide_data.get("note_to_send_to_client", "")
        guide_path  = build_client_guide_docx(guide_data, client_name)
        safe        = re.sub(r'[^\w\s-]', '', guide_name).strip().replace(' ', '_')

        # Build the Slack comment: note to send + Trainerize habits for the coach
        habits      = guide_data.get("trainerize_habits", [])
        habits_text = ""
        if habits:
            habits_list = "\n".join(f"  - {h}" for h in habits)
            habits_text = (
                f"\n\n*Trainerize Habits to Add (for any client using this guide):*\n"
                f"{habits_list}\n"
                f"_Add whichever the client decides to prioritize as a trackable habit in the app._"
            )

        upload_file_to_slack(channel_id, thread_ts, guide_path, f"{safe}.docx",
                             f"*Note to send to client:*\n{note}{habits_text}")
        print(f"Client guide posted for {client_name}")

        # 2. Market research
        print("Step 2: market research...")
        search_results = run_guide_market_research(topic)

        # 3. Selling opportunities
        print("Step 3: selling opportunities...")
        selling_data  = generate_selling_opportunities_content(guide_data, topic, search_results)
        selling_path  = build_selling_opportunities_docx(selling_data, guide_name)

        # 4. Digital product
        print("Step 4: digital product...")
        digital_data  = generate_digital_product_content(guide_data, selling_data,
                                                         transcript, fathom_summary)
        digital_path  = build_digital_product_docx(digital_data)
        digital_name  = re.sub(r'[^\w\s-]', '', digital_data.get("product_name", guide_name)).strip().replace(' ', '_')

        # 5. Lead magnet
        print("Step 5: lead magnet...")
        lm_data  = generate_lead_magnet_content(guide_data, selling_data,
                                                transcript, fathom_summary)
        lm_path  = build_lead_magnet_docx(lm_data)
        lm_name  = re.sub(r'[^\w\s-]', '', lm_data.get("freebie_name", guide_name)).strip().replace(' ', '_')

        # 6. Post 3 docs to #projects-launches-summaries
        if projects_channel_id:
            keyword    = selling_data.get("automation_keyword", guide_data.get("keyword", ""))
            intro_text = (
                f"*Guide Packet: {guide_name}*\n"
                f"From {client_name}'s call | Topic: {topic}\n"
                f"Automation keyword: `{keyword}`\n"
                f"3 documents attached for review:"
            )
            hdr = {"Authorization": f"Bearer {SLACK_BOT_TOKEN}"}
            ir  = requests.post("https://slack.com/api/chat.postMessage",
                                headers=hdr,
                                json={"channel": projects_channel_id, "text": intro_text})
            its = ir.json().get("ts") if ir.json().get("ok") else None

            upload_file_to_slack(projects_channel_id, its, selling_path,
                                 f"{safe}_Selling_Opportunities.docx",
                                 "Selling Opportunities Assessment")
            upload_file_to_slack(projects_channel_id, its, digital_path,
                                 f"{digital_name}_Digital_Product.docx",
                                 "Digital Product (Full Version)")
            upload_file_to_slack(projects_channel_id, its, lm_path,
                                 f"{lm_name}_Lead_Magnet.docx",
                                 "Lead Magnet / Freebie Version")
            print("All 3 docs posted to #projects-launches-summaries")

        print(f"=== Guide pipeline complete: {client_name} | {guide_name} ===")

    except Exception as e:
        print(f"Guide pipeline error for {client_name}: {e}")
        import traceback
        traceback.print_exc()
    finally:
        for p in [guide_path, selling_path, digital_path, lm_path]:
            if p and os.path.exists(p):
                try:
                    os.remove(p)
                except Exception:
                    pass
